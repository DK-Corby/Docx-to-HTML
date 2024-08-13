import sys
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

def convert_html_to_docx(docx_path, output_path):
    # Load the document
    doc = Document(docx_path)
    
    # Iterate through all paragraphs and runs
    for paragraph in doc.paragraphs:
        new_runs = []
        for run in paragraph.runs:
            # Use BeautifulSoup to parse the run text
            soup = BeautifulSoup(run.text, 'html.parser')
            
            for elem in soup:
                if elem.name == 'b':
                    bold_run = paragraph.add_run(elem.get_text())
                    bold_run.bold = True
                    new_runs.append(bold_run)
                elif elem.name == 'i':
                    italic_run = paragraph.add_run(elem.get_text())
                    italic_run.italic = True
                    new_runs.append(italic_run)
                else:
                    normal_run = paragraph.add_run(elem if isinstance(elem, str) else elem.get_text())
                    new_runs.append(normal_run)
            run.clear()

        # Remove old runs and add new runs to the paragraph
        for run in paragraph.runs:
            paragraph._element.remove(run._element)
        for new_run in new_runs:
            paragraph.add_run(new_run.text, style=new_run.style)
            if new_run.bold:
                paragraph.runs[-1].bold = True
            if new_run.italic:
                paragraph.runs[-1].italic = True

    # Save the modified document
    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_to_native.py <input_docx_path> <output_docx_path>")
    else:
        input_docx_path = sys.argv[1]
        output_docx_path = sys.argv[2]
        convert_html_to_docx(input_docx_path, output_docx_path)

